from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

def set_font(run, size=12, bold=False, color=None, font_name='Iskoola Pota'):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1, color=(26, 8, 0)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, size=14 if level==1 else 12, bold=True, color=color)
    if level == 1:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        # Add bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'C8960C')
        pBdr.append(bottom)
        pPr.append(pBdr)
        # Shading
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '2A1200')
        pPr.append(shd)
        run.font.color.rgb = RGBColor(255, 255, 255)
    return p

def add_para(doc, text, size=11, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, color=(26,8,0)):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_table(doc, headers, rows):
    num_cols = len(headers)
    num_rows = len(rows)
    table = doc.add_table(rows=1+num_rows, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=11, bold=True, color=(255,255,255))
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '2A1200')
        tcPr.append(shd)
    # Data rows
    for ri in range(num_rows):
        row_data = rows[ri]
        row = table.rows[ri+1]
        is_last = (ri == num_rows-1)
        for ci in range(num_cols):
            val = row_data[ci] if ci < len(row_data) else ''
            cell = row.cells[ci]
            cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            set_font(run, size=11, bold=is_last, color=(26,8,0))
            if ri % 2 == 1:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'FDF6ED')
                tcPr.append(shd)
    doc.add_paragraph()

# ==================== COVER PAGE ====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
try:
    run = p.add_run()
    run.add_picture('Logo.png', width=Inches(1.5))
except:
    run = p.add_run('ALM SK Ceylon Cashews')
    set_font(run, size=14, bold=True, color=(200,150,12))

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ALM SK Ceylon Cashews (Pvt) Ltd')
set_font(run, size=20, bold=True, color=(26,8,0))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('"Pure Ceylon රසය"')
set_font(run, size=13, bold=False, color=(200,150,12))

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ව්‍යාපාර සැලැස්ම')
set_font(run, size=18, bold=True, color=(26,8,0))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ණය අයදුම්පත — Bank of Ceylon')
set_font(run, size=13, color=(100,50,0))

doc.add_paragraph()

# Cover info table
table = doc.add_table(rows=10, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
info = [
    ('ව්‍යාපාරයේ නම', 'ALM SK Ceylon Cashews (Pvt) Ltd'),
    ('ව්‍යාපාර වර්ගය', 'පවුලේ හවුල් ව්‍යාපාරය — කජු සැකසීම හා වෙළඳාම'),
    ('හවුල්කරුවන්', ''),
    ('   1.', 'Lakruwan Somarathna'),
    ('   2.', 'Anupama Somarathna'),
    ('   3.', 'Manoja Somarathna'),
    ('ණය අයදුම්කරු', 'Anupama Somarathna'),
    ('ලිපිනය', 'Diulwewa, Anamaduwa, Kurunegala, Sri Lanka'),
    ('දුරකථන', '+94 78 381 9650'),
    ('දිනය', 'අප්‍රේල් 2026'),
]
for i, (k, v) in enumerate(info):
    row = table.rows[i]
    run_k = row.cells[0].paragraphs[0].add_run(k)
    set_font(run_k, size=11, bold=True, color=(26,8,0))
    run_v = row.cells[1].paragraphs[0].add_run(v)
    set_font(run_v, size=11, color=(26,8,0))

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ඉල්ලා සිටින ණය මුදල: රු. 5,000,000 (රුපියල් ලක්ෂ පනහක්)')
set_font(run, size=12, bold=True, color=(26,8,0))

doc.add_page_break()

# ==================== SECTION 1 ====================
add_heading(doc, '1. සාරාංශය (Executive Summary)')
add_para(doc, 'ALM SK Ceylon Cashews (Pvt) Ltd යනු 2026 දී ආරම්භ කරන ලද, Diulwewa, Anamaduwa, Kurunegala හි පිහිටි පවුලේ හවුල් ව්‍යාපාරයකි. Somarathna පවුලේ දරුවන් තිදෙනා — Lakruwan, Anupama, සහ Manoja Somarathna — විසින් ශ්‍රී ලංකාවේ ප්‍රිමියම් කජු නිෂ්පාදනය හා අලෙවිය සඳහා මෙම ව්‍යාපාරය ආරම්භ කරන ලදී.')
add_para(doc, 'ව්‍යාපාරයේ ප්‍රධාන නිෂ්පාදන: අමු කජු, බදින ලද කජු, සහ රස කළ කජු — දේශීය හා ජාත්‍යන්තර වෙළඳපොළ ඉලක්ක කරගනිමින්.')
doc.add_paragraph()

add_heading(doc, '   ණය ඉල්ලීමේ සාරාංශය', level=2)
add_table(doc,
    ['විස්තරය', 'මුදල (රු.)'],
    [
        ('ඉල්ලා සිටින ණය මුදල', '5,000,000'),
        ('කජු තොග මිලදී ගැනීම', '4,000,000'),
        ('හදිසි / ක්‍රියාකාරී ප්‍රාග්ධනය', '1,000,000'),
        ('මුළු එකතුව', '5,000,000'),
    ]
)

# ==================== SECTION 2 ====================
add_heading(doc, '2. ව්‍යාපාර විස්තරය')
add_heading(doc, '   2.1 සමාගම් විස්තර', level=2)
add_table(doc,
    ['විස්තරය', 'තොරතුරු'],
    [
        ('ව්‍යාපාරයේ නම', 'ALM SK Ceylon Cashews (Pvt) Ltd'),
        ('ව්‍යාපාර වර්ගය', 'පවුලේ හවුල් ව්‍යාපාරය'),
        ('හවුල්කරු 01', 'Lakruwan Somarathna'),
        ('හවුල්කරු 02 (ණය අයදුම්කරු)', 'Anupama Somarathna'),
        ('හවුල්කරු 03', 'Manoja Somarathna'),
        ('ආරම්භ කළ වර්ෂය', '2026'),
        ('ලිපිනය', 'Diulwewa, Anamaduwa, Kurunegala'),
        ('දුරකථන', '+94 78 381 9650'),
        ('විද්‍යුත් තැපෑල', 'almskceyloncashews@gmail.com'),
        ('වෙබ් අඩවිය', 'almskceyloncashews.netlify.app'),
    ]
)

add_heading(doc, '   2.2 නිෂ්පාදන හා සේවා', level=2)
add_table(doc,
    ['නිෂ්පාදනය', 'විස්තරය', 'ඉලක්ක වෙළඳපොළ'],
    [
        ('අමු කජු', '100% ස්වාභාවික, W180–W450 අපනයන ශ්‍රේණිය', 'තොග ගැනුම්කරුවන්, අපනයන'),
        ('බදින ලද කජු', 'සාම්ප්‍රදායික ක්‍රමයෙන් රන්වන් ලෙස බදින ලද', 'සිල්ලර, දේශීය වෙළඳපොළ'),
        ('රස කළ කජු', 'කුළු බඩු, මී පැණි, ලුණු බටර් රස', 'සිල්ලර, තෑගි පැකේජ'),
    ]
)

doc.add_page_break()

# ==================== SECTION 3 ====================
add_heading(doc, '3. වෙළඳපොළ විශ්ලේෂණය')
add_heading(doc, '   3.1 ඉලක්ක වෙළඳපොළ', level=2)
add_table(doc,
    ['කොටස', 'විස්තරය'],
    [
        ('දේශීය සිල්ලර', 'සුපිරි වෙළඳසැල්, කඩ, සෘජු පාරිභෝගිකයන්'),
        ('තොග ගැනුම්කරුවන්', 'හෝටල්, අවන්හල්, ආහාර නිෂ්පාදකයන්'),
        ('අපනයන වෙළඳපොළ', 'ජාත්‍යන්තර ආනයනකරුවන් — මැදපෙරදිග, යුරෝපය'),
        ('B2B සැපයුම්කරුවන්', 'ශ්‍රී ලංකාව පුරා අමු කජු ගොවීන්'),
    ]
)

add_heading(doc, '   3.2 තරඟකාරී වාසි', level=2)
for item in [
    '✔ ගොවීන් සමඟ සෘජු සම්බන්ධතා — අඩු ලබා ගැනීමේ පිරිවැය',
    '✔ සාම්ප්‍රදායික ක්‍රම භාවිතයෙන් ප්‍රිමියම් ගුණාත්මක සැකසීම',
    '✔ වෙබ් අඩවිය හා ඩිජිටල් අලෙවිකරණය හරහා ස්ථාපිත මාර්ගගත පැවැත්ම',
    '✔ කුරුණෑගල — ශ්‍රී ලංකාවේ ප්‍රධාන කජු වගා කලාපයේ පිහිටීම',
    '✔ සිල්ලර හා තොග වෙළඳපොළ සඳහා නම්‍යශීලී නිෂ්පාදන පරාසය',
]:
    add_para(doc, item, size=11)

doc.add_paragraph()

# ==================== SECTION 4 ====================
add_heading(doc, '4. ණය භාවිත සැලැස්ම')
add_para(doc, 'ඉල්ලා සිටින රු. 5,000,000 (ලක්ෂ පනහ) ණය මුදල පහත පරිදි භාවිත කිරීමට සැලසුම් කර ඇත:')
add_table(doc,
    ['අරමුණ', 'මුදල (රු.)', 'ප්‍රතිශතය'],
    [
        ('අමු කජු තොග මිලදී ගැනීම', '4,000,000', '80%'),
        ('හදිසි / ක්‍රියාකාරී ප්‍රාග්ධනය', '1,000,000', '20%'),
        ('මුළු එකතුව', '5,000,000', '100%'),
    ]
)
add_para(doc, 'තොග ආයෝජන උපාය: රු. 4,000,000 ගොවීන්ගෙන් සෘජුව ගොවිපල-දොරකඩ මිලට අමු කජු තොග මිලදී ගැනීමට භාවිත කෙරේ. තොග මිලදී ගැනීම kg එකකට පිරිවැය සැලකිය යුතු ලෙස අඩු කරයි.')
add_para(doc, 'හදිසි සංචිතය: රු. 1,000,000 ක්‍රියාකාරී වියදම්, ඇසුරුම්, ප්‍රවාහනය සහ අනපේක්ෂිත පිරිවැය සඳහා ක්‍රියාකාරී ප්‍රාග්ධනයක් ලෙස පවත්වා ගනු ලැබේ.')

doc.add_page_break()

# ==================== SECTION 5 ====================
add_heading(doc, '5. මූල්‍ය ප්‍රක්ෂේපණ')
add_heading(doc, '   5.1 ආදායම් ප්‍රක්ෂේපණය (1 වන වර්ෂය)', level=2)
add_table(doc,
    ['මාසය', 'ඇස්තමේන්තු විකුණුම් (රු.)', 'ඇස්තමේන්තු ලාභය (රු.)'],
    [
        ('මාස 1–2', '300,000', '60,000'),
        ('මාස 3–4', '500,000', '100,000'),
        ('මාස 5–6', '700,000', '140,000'),
        ('මාස 7–9', '900,000', '180,000'),
        ('මාස 10–12', '1,200,000', '240,000'),
        ('1 වන වර්ෂ මුළු එකතුව', '3,600,000', '720,000'),
    ]
)

add_heading(doc, '   5.2 ණය ආපසු ගෙවීමේ හැකියාව', level=2)
add_table(doc,
    ['විස්තරය', 'මුදල (රු.)'],
    [
        ('ණය මුදල', '5,000,000'),
        ('ඇස්තමේන්තු ණය කාලය', 'වර්ෂ 3–5'),
        ('ඇස්තමේන්තු මාසික ගෙවීම', '70,000'),
        ('ප්‍රක්ෂේපිත මාසික ආදායම (සාමාන්‍ය)', '300,000 – 1,200,000'),
        ('ආපසු ගෙවීමේ හැකියාව', '✔ හැකි'),
    ]
)

# ==================== SECTION 6 ====================
add_heading(doc, '6. ව්‍යාපාර උපාය හා වර්ධන සැලැස්ම')
add_heading(doc, '   6.1 කෙටි කාලීන (1 වන වර්ෂය)', level=2)
for item in ['✔ ණය මුදල් භාවිතයෙන් අමු කජු තොග මිලදී ගැනීම', '✔ කුරුණෑගල දිස්ත්‍රික්කයේ සිල්ලර බෙදාහැරීම ස්ථාපිත කිරීම', '✔ වෙබ් අඩවිය හා WhatsApp හරහා මාර්ගගත විකිණීම', '✔ Alibaba, TradeIndia හරහා අපනයන වේදිකාවල ලියාපදිංචිය']:
    add_para(doc, item, size=11)

add_heading(doc, '   6.2 මධ්‍යම කාලීන (වර්ෂ 2–3)', level=2)
for item in ['✔ කොළඹ හා අනෙකුත් ප්‍රධාන නගරවලට ව්‍යාප්ත කිරීම', '✔ මැදපෙරදිග හෝ යුරෝපයට පළමු අපනයන ඇණවුම', '✔ නව රස සමඟ නිෂ්පාදන පරාසය වැඩි කිරීම', '✔ ව්‍යාපාර ආදායමෙන් ණය ආපසු ගෙවීම ආරම්භ කිරීම']:
    add_para(doc, item, size=11)

add_heading(doc, '   6.3 දිගු කාලීන (වර්ෂ 4–5)', level=2)
for item in ['✔ තමන්ගේම සැකසුම් පහසුකම් ස්ථාපිත කිරීම', '✔ අපනයන සහතික (SLS, ISO) ලබා ගැනීම', '✔ ණය සම්පූර්ණයෙන් ආපසු ගෙවීම', '✔ ස්වාධීනව ජාත්‍යන්තර වෙළඳපොළ ව්‍යාප්ත කිරීම']:
    add_para(doc, item, size=11)

# ==================== SECTION 7 ====================
add_heading(doc, '7. ප්‍රකාශය හා අත්සන')
doc.add_paragraph()
add_para(doc, 'අපි, ALM SK Ceylon Cashews (Pvt) Ltd හි හවුල්කරුවන් වශයෙන්, මෙම ව්‍යාපාර සැලැස්මේ සඳහන් සියලු තොරතුරු අපගේ දැනුමට හා විශ්වාසයට අනුව සත්‍ය, නිවැරදි හා සම්පූර්ණ බව ප්‍රකාශ කරමු.')
doc.add_paragraph()
add_para(doc, 'රු. 5,000,000 (ලක්ෂ පනහ) ණය මුදල Anupama Somarathna විසින් ප්‍රාථමික අයදුම්කරු ලෙස ඉල්ලා සිටී. හවුල්කරුවන් තිදෙනාම ව්‍යාපාරයට හා එහි බැඳීම්වලට ඒකාබද්ධව කැපවී සිටිති.')
doc.add_paragraph()
doc.add_paragraph()

# Signature table
sig_table = doc.add_table(rows=2, cols=3)
sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
labels = ['හවුල්කරු 01', 'හවුල්කරු 02\n(ණය අයදුම්කරු)', 'හවුල්කරු 03']
names = ['Lakruwan Somarathna', 'Anupama Somarathna', 'Manoja Somarathna']
for i in range(3):
    cell = sig_table.rows[0].cells[i]
    run = cell.paragraphs[0].add_run(labels[i])
    set_font(run, size=11, bold=True)
    cell2 = sig_table.rows[1].cells[i]
    run2 = cell2.paragraphs[0].add_run(names[i] + '\n\nඅත්සන: ________________\n\nදිනය: ________________')
    set_font(run2, size=11)

doc.add_paragraph()
doc.add_paragraph()

bank_table = doc.add_table(rows=2, cols=2)
bank_table.alignment = WD_TABLE_ALIGNMENT.CENTER
bank_labels = ['බැංකු නිලධාරී අත්සන', 'ශාඛා මුද්‍රාව']
bank_content = ['නම: ________________\nතනතුර: ________________\nදිනය: ________________', '']
for i in range(2):
    cell = bank_table.rows[0].cells[i]
    run = cell.paragraphs[0].add_run(bank_labels[i])
    set_font(run, size=11, bold=True)
    cell2 = bank_table.rows[1].cells[i]
    run2 = cell2.paragraphs[0].add_run(bank_content[i])
    set_font(run2, size=11)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ALM SK Ceylon Cashews (Pvt) Ltd · Diulwewa, Anamaduwa, Kurunegala, Sri Lanka\n+94 78 381 9650 · almskceyloncashews@gmail.com · almskceyloncashews.netlify.app\n"Pure Ceylon රසය"')
set_font(run, size=10, color=(100,50,0))

# Save
output = 'Business Plan Sinhala - ALM SK Ceylon Cashews.docx'
doc.save(output)
print(f'Document saved: {output}')
