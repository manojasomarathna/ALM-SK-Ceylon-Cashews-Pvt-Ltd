from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

def set_font(run, size=12, bold=False, color=None, font_name='Calibri'):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, size=14 if level==1 else 12, bold=True, color=(255,255,255) if level==1 else (26,8,0))
    if level == 1:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '2A1200')
        pPr.append(shd)
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'C8960C')
        pBdr.append(bottom)
        pPr.append(pBdr)
    else:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'C8960C')
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p

def add_para(doc, text, size=11, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, color=(26,8,0)):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_table(doc, headers, rows):
    num_cols = len(headers)
    num_rows = len(rows)
    table = doc.add_table(rows=1+num_rows, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=11, bold=True, color=(255,255,255))
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '2A1200')
        tcPr.append(shd)
    for ri in range(num_rows):
        row_data = rows[ri]
        row = table.rows[ri+1]
        is_last = (ri == num_rows-1)
        for ci in range(num_cols):
            val = row_data[ci] if ci < len(row_data) else ''
            cell = row.cells[ci]
            cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            set_font(run, size=11, bold=is_last, color=(26,8,0))
            if ri % 2 == 1:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'FDF6ED')
                tcPr.append(shd)
    doc.add_paragraph()

# ==================== COVER PAGE ====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
try:
    run = p.add_run()
    run.add_picture('Logo.png', width=Inches(1.5))
except:
    run = p.add_run('ALM SK Ceylon Cashews')
    set_font(run, size=14, bold=True, color=(200,150,12))

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ALM SK Ceylon Cashews (Pvt) Ltd')
set_font(run, size=22, bold=True, color=(26,8,0))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('"Taste of Pure Ceylon"')
set_font(run, size=13, color=(200,150,12))

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('BUSINESS PLAN')
set_font(run, size=20, bold=True, color=(26,8,0))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Loan Application — Bank of Ceylon')
set_font(run, size=13, color=(100,50,0))

doc.add_paragraph()

table = doc.add_table(rows=10, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
info = [
    ('Business Name', 'ALM SK Ceylon Cashews (Pvt) Ltd'),
    ('Business Type', 'Family Partnership — Cashew Processing & Trading'),
    ('Partners', ''),
    ('   1.', 'Lakruwan Somarathna'),
    ('   2.', 'Anupama Somarathna'),
    ('   3.', 'Manoja Somarathna'),
    ('Loan Applicant', 'Mrs. Anupama Somarathna'),
    ('Address', 'Diulwewa, Anamaduwa, Kurunegala, Sri Lanka'),
    ('Phone', '+94 78 381 9650'),
    ('Date', 'April 2026'),
]
for i, (k, v) in enumerate(info):
    row = table.rows[i]
    run_k = row.cells[0].paragraphs[0].add_run(k)
    set_font(run_k, size=11, bold=True, color=(26,8,0))
    run_v = row.cells[1].paragraphs[0].add_run(v)
    set_font(run_v, size=11, color=(26,8,0))

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Loan Amount Requested: Rs. 5,000,000 (Fifty Lakhs)')
set_font(run, size=12, bold=True, color=(26,8,0))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('STRICTLY CONFIDENTIAL')
set_font(run, size=11, bold=True, color=(255,255,255))
pPr = p._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:val'), 'clear')
shd.set(qn('w:color'), 'auto')
shd.set(qn('w:fill'), '1A0800')
pPr.append(shd)

doc.add_page_break()

# ==================== SECTION 1 ====================
add_heading(doc, '1. Executive Summary')
add_para(doc, 'ALM SK Ceylon Cashews (Pvt) Ltd is a family-owned cashew processing and trading business established in 2026, located in Diulwewa, Anamaduwa, Kurunegala, Sri Lanka. The business is founded by the Somarathna family — Lakruwan, Anupama, and Manoja Somarathna — with a deep-rooted passion for Sri Lankan cashew farming.')
add_para(doc, 'The loan application is submitted by Mrs. Anupama Somarathna as the primary applicant, supported by her employment income, while all three partners are equally committed to the success of the business.')
doc.add_paragraph()

add_heading(doc, '   Loan Summary', level=2)
add_table(doc,
    ['Description', 'Amount (Rs.)'],
    [
        ('Loan Amount Requested', '5,000,000'),
        ('Raw Cashew Stock Purchase', '4,000,000'),
        ('Emergency / Working Capital', '1,000,000'),
        ('Total', '5,000,000'),
    ]
)

# ==================== SECTION 2 ====================
add_heading(doc, '2. Business Overview')
add_heading(doc, '   2.1 Company Details', level=2)
add_table(doc,
    ['Detail', 'Information'],
    [
        ('Business Name', 'ALM SK Ceylon Cashews (Pvt) Ltd'),
        ('Business Type', 'Family Partnership'),
        ('Partner 01', 'Lakruwan Somarathna'),
        ('Partner 02 (Loan Applicant)', 'Anupama Somarathna'),
        ('Partner 03', 'Manoja Somarathna'),
        ('Year Established', '2026'),
        ('Address', 'Diulwewa, Anamaduwa, Kurunegala'),
        ('Phone', '+94 78 381 9650'),
        ('Email', 'almskceyloncashews@gmail.com'),
        ('Website', 'almskceyloncashews.netlify.app'),
    ]
)

add_heading(doc, '   2.2 Products & Services', level=2)
add_table(doc,
    ['Product', 'Description', 'Target Market'],
    [
        ('Raw Cashews', '100% natural, export-grade W180–W450', 'Bulk buyers, Export'),
        ('Roasted Cashews', 'Slow-roasted, lightly salted, premium packaging', 'Retail, Local market'),
        ('Flavored Cashews', 'Spicy, honey glazed, salted butter varieties', 'Retail, Gift packs'),
    ]
)

doc.add_page_break()

# ==================== SECTION 3 ====================
add_heading(doc, '3. Market Analysis')
add_heading(doc, '   3.1 Target Market', level=2)
add_table(doc,
    ['Segment', 'Description'],
    [
        ('Local Retail', 'Supermarkets, grocery stores, direct consumers'),
        ('Bulk Buyers', 'Hotels, restaurants, food manufacturers'),
        ('Export Market', 'International importers — Middle East, Europe'),
        ('B2B Suppliers', 'Raw cashew farmers across Sri Lanka'),
    ]
)

add_heading(doc, '   3.2 Competitive Advantages', level=2)
for item in [
    '✔ Direct relationships with farmers — lower procurement cost',
    '✔ Premium quality processing using traditional methods',
    '✔ Established online presence via website and digital marketing',
    '✔ Located in Kurunegala — prime cashew growing region of Sri Lanka',
    '✔ Flexible product range catering to retail and bulk markets',
]:
    add_para(doc, item, size=11)

doc.add_paragraph()

# ==================== SECTION 4 ====================
add_heading(doc, '4. Loan Utilization Plan')
add_para(doc, 'The requested loan of Rs. 5,000,000 (Fifty Lakhs) will be utilized as follows:')
add_table(doc,
    ['Purpose', 'Amount (Rs.)', 'Percentage'],
    [
        ('Raw Cashew Stock Purchase', '4,000,000', '80%'),
        ('Emergency / Working Capital Reserve', '1,000,000', '20%'),
        ('Total', '5,000,000', '100%'),
    ]
)
add_para(doc, 'Stock Investment Strategy: Rs. 4,000,000 will be used to purchase raw cashew stock in bulk directly from farmers at farm-gate prices, significantly reducing per-kg cost and allowing higher profit margins.')
add_para(doc, 'Emergency Reserve: Rs. 1,000,000 will be maintained as working capital to cover operational expenses, packaging, transportation, and unforeseen costs.')

doc.add_page_break()

# ==================== SECTION 5 ====================
add_heading(doc, '5. Financial Projections')
add_heading(doc, '   5.1 Revenue Projection (Year 1)', level=2)
add_table(doc,
    ['Month', 'Estimated Sales (Rs.)', 'Estimated Profit (Rs.)'],
    [
        ('Month 1–2', '300,000', '60,000'),
        ('Month 3–4', '500,000', '100,000'),
        ('Month 5–6', '700,000', '140,000'),
        ('Month 7–9', '900,000', '180,000'),
        ('Month 10–12', '1,200,000', '240,000'),
        ('Year 1 Total', '3,600,000', '720,000'),
    ]
)

add_heading(doc, '   5.2 Loan Repayment Capacity', level=2)
add_table(doc,
    ['Detail', 'Amount (Rs.)'],
    [
        ('Loan Amount', '5,000,000'),
        ('Estimated Loan Period', '3–5 Years'),
        ('Estimated Monthly Repayment', '70,000'),
        ('Projected Monthly Revenue (Avg)', '300,000 – 1,200,000'),
        ('Repayment Feasibility', '✔ Feasible'),
    ]
)

# ==================== SECTION 6 ====================
add_heading(doc, '6. Business Strategy & Growth Plan')
add_heading(doc, '   6.1 Short Term (Year 1)', level=2)
for item in ['✔ Purchase bulk raw cashew stock using loan funds', '✔ Establish retail distribution in Kurunegala district', '✔ Launch online sales via website and WhatsApp', '✔ Register on export platforms (Alibaba, TradeIndia)']:
    add_para(doc, item, size=11)

add_heading(doc, '   6.2 Medium Term (Year 2–3)', level=2)
for item in ['✔ Expand to Colombo and other major cities', '✔ Secure first export order to Middle East or Europe', '✔ Increase product range with new flavors', '✔ Begin loan repayment from business revenue']:
    add_para(doc, item, size=11)

add_heading(doc, '   6.3 Long Term (Year 4–5)', level=2)
for item in ['✔ Establish own processing facility', '✔ Obtain export certifications (SLS, ISO)', '✔ Full loan repayment', '✔ Expand to international markets independently']:
    add_para(doc, item, size=11)

# ==================== SECTION 7 ====================
add_heading(doc, '7. Declaration & Signatures')
doc.add_paragraph()
add_para(doc, 'We, the undersigned partners of ALM SK Ceylon Cashews (Pvt) Ltd, hereby declare that all information provided in this business plan is true, accurate, and complete to the best of our knowledge.')
doc.add_paragraph()
add_para(doc, 'The loan of Rs. 5,000,000 (Fifty Lakhs) is applied for by Mrs. Anupama Somarathna as the primary applicant. All three partners are jointly committed to the business and its obligations.')
doc.add_paragraph()
doc.add_paragraph()

sig_table = doc.add_table(rows=2, cols=3)
sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
labels = ['Partner 01', 'Partner 02\n(Loan Applicant)', 'Partner 03']
names = ['Lakruwan Somarathna', 'Anupama Somarathna', 'Manoja Somarathna']
for i in range(3):
    cell = sig_table.rows[0].cells[i]
    run = cell.paragraphs[0].add_run(labels[i])
    set_font(run, size=11, bold=True)
    cell2 = sig_table.rows[1].cells[i]
    run2 = cell2.paragraphs[0].add_run(names[i] + '\n\nSignature: ________________\n\nDate: ________________')
    set_font(run2, size=11)

doc.add_paragraph()
doc.add_paragraph()

bank_table = doc.add_table(rows=2, cols=2)
bank_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (lbl, val) in enumerate([('Bank Officer Signature', 'Name: ________________\nDesignation: ________________\nDate: ________________'), ('Branch Stamp', '')]):
    cell = bank_table.rows[0].cells[i]
    run = cell.paragraphs[0].add_run(lbl)
    set_font(run, size=11, bold=True)
    cell2 = bank_table.rows[1].cells[i]
    run2 = cell2.paragraphs[0].add_run(val)
    set_font(run2, size=11)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ALM SK Ceylon Cashews (Pvt) Ltd · Diulwewa, Anamaduwa, Kurunegala, Sri Lanka\n+94 78 381 9650 · almskceyloncashews@gmail.com · almskceyloncashews.netlify.app\n"Taste of Pure Ceylon"')
set_font(run, size=10, color=(100,50,0))

output = 'Business Plan English - ALM SK Ceylon Cashews.docx'
doc.save(output)
print(f'Saved: {output}')
